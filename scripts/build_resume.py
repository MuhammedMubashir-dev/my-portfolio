import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
PUBLIC = ROOT / "public"

INK = RGBColor(33, 33, 33)
MUTED = RGBColor(100, 100, 100)
HEADING = RGBColor(26, 58, 95)
LINK = RGBColor(26, 58, 95)
BORDER = "B0B0B0"

ROLE_START = date(2026, 4, 1)
PROMOTION_DATE = date(2026, 7, 1)

VARIANTS = {
    "react": {
        "file": "muhammed-mubashir-k-resume-react.docx",
        "subtitle": "React Developer  |  React.js · Next.js · React Native · TypeScript",
        "summary": (
            "React Developer with production experience shipping web and mobile applications using "
            "React.js, Next.js, and React Native. Built a multi-tenant Next.js e-commerce platform with "
            "8+ theme packs, delivered bilingual (English/Arabic) storefronts with RTL support, and "
            "developed a React Native business networking app. Skilled in REST API integration, SSR/ISR "
            "performance optimization, payment gateway integration, and responsive UI development. "
            "Also experienced in Flutter for cross-platform mobile development."
        ),
        "current_bullets": [
            "Engineered a multi-tenant Next.js e-commerce platform with 8+ production theme packs, tenant onboarding, and KV-based tenant switching serving live client storefronts.",
            "Built live search autocomplete with debounced requests and request cancellation, integrated across multiple theme navigation bars with Algolia-powered product discovery.",
            "Developed a React Native business networking app with company profile management, connection request workflows, subscription handling, and deep link integration.",
            "Delivered bilingual (English/Arabic) storefronts with full RTL layout support, i18n translation systems, ISR caching, and dynamic SEO metadata with sitemap generation.",
            "Integrated Razorpay payment gateway across tenant checkouts and mobile apps; built OTP authentication, guest checkout, and COD flows with WhatsApp order notifications.",
            "Also built Flutter mobile applications for POS operations and logistics, covering billing, inventory management, and delivery workflows.",
        ],
        "trainee_bullets": [
            "Integrated REST APIs for authentication, product catalogs, orders, and checkout across Next.js storefronts and Flutter mobile applications.",
            "Built dynamic category-driven property filter architecture with multi-select filtering and a review system supporting multi-image uploads.",
            "Implemented Google Maps navigation, infinite-scroll pagination, and date-based filtering for logistics delivery management screens.",
        ],
        "projects": [
            (
                "FUNZCART / CloudPOS Web",
                "Multi-Tenant E-Commerce Platform (Next.js)",
                "Built and maintained a multi-tenant Next.js platform with 8+ production theme packs "
                "(Aurora, Nova, Toy, Silk, Ornament, Harvest) and tenant storefronts including Luzine "
                "Bakes. Designed dynamic property filter architecture, live search autocomplete, and a "
                "review system with multi-image uploads. Delivered Luzine Bakes with guest checkout, "
                "inline OTP verification, and a COD flow with WhatsApp order notifications.",
            ),
            (
                "Juice World",
                "Production Storefront (Next.js · i18n · ISR)",
                "Developed a production storefront with API-driven banners, products, services, and "
                "testimonials. Implemented English/Arabic i18n with full RTL support, ISR caching, "
                "dynamic SEO metadata, and sitemap generation. Refactored to server component "
                "architecture for production readiness.",
            ),
            (
                "Connect App",
                "Business Networking Application (React Native)",
                "Built a business networking app enabling companies to create and share professional "
                "profiles, manage connection requests, and handle subscriptions. Implemented deep "
                "linking, profile sharing via shortened URLs, group navigation, and username auto-sync.",
            ),
            (
                "Ganvin",
                "Logistics & Delivery Ecosystem (Flutter)",
                "Built Executive and Customer apps for the Ganvin logistics ecosystem. Implemented "
                "Google Maps navigation, infinite-scroll pagination, Razorpay payments, and resolved "
                "timezone/API issues with the backend team.",
            ),
            (
                "EPOSMOB",
                "Mobile POS & Inventory Management (Flutter)",
                "Developed a mobile POS application covering billing, Day Close workflows, bilingual "
                "receipt printing, stock reports, purchase orders, and supplier/customer management.",
            ),
        ],
        "skills": [
            ("React Ecosystem", "React.js, Next.js, React Native, TypeScript, JavaScript"),
            ("Web Architecture", "SSR, ISR, Dynamic Routing, SEO, Sitemap Generation, i18n, RTL Support, Responsive Design"),
            ("Frontend", "HTML5, CSS3, Tailwind CSS, State Management"),
            ("Integrations", "Razorpay, Algolia Search, Google Maps, WhatsApp Notifications, Barcode Generation, Deep Linking"),
            ("Commerce & Multi-Tenant", "Multi-Tenant Architecture (8+ themes), Tenant Onboarding, Product Search, Checkout Flows, Order Management"),
            ("Also experienced in", "Flutter, Dart, Python, Django, MySQL, SQLite"),
            ("Tools", "Git, GitHub, VS Code, Postman, Figma"),
        ],
    },
    "react-native": {
        "file": "muhammed-mubashir-k-resume-react-native.docx",
        "subtitle": "React Native Developer  |  React Native · React.js · Next.js · TypeScript",
        "summary": (
            "React Native Developer with production experience building cross-platform mobile "
            "applications. Shipped a business networking app with deep linking, profile sharing, "
            "subscription management, and group navigation. Also built production web applications "
            "using React.js and Next.js, including a multi-tenant e-commerce platform with 8+ theme "
            "packs and bilingual (English/Arabic) storefronts with RTL support. "
            "Additionally experienced in Flutter for mobile development."
        ),
        "current_bullets": [
            "Developed a React Native business networking app with company profile management, connection request workflows, subscription cancellation/reactivation, and deep link integration.",
            "Implemented profile sharing via shortened URLs, group navigation, username auto-sync, and exit confirmation flows in the React Native application.",
            "Engineered a multi-tenant Next.js e-commerce platform with 8+ production theme packs, tenant onboarding, and KV-based tenant switching serving live client storefronts.",
            "Built live search autocomplete with debounced requests and request cancellation, integrated across multiple theme navigation bars with Algolia-powered product discovery.",
            "Delivered bilingual (English/Arabic) storefronts with full RTL layout support, i18n translation systems, ISR caching, and dynamic SEO metadata.",
            "Also built Flutter mobile applications for POS operations and logistics, covering billing, inventory management, and delivery workflows.",
        ],
        "trainee_bullets": [
            "Integrated REST APIs for authentication, product catalogs, orders, and checkout across Next.js storefronts and Flutter mobile applications.",
            "Built dynamic category-driven property filter architecture with multi-select filtering and a review system supporting multi-image uploads.",
            "Implemented Google Maps navigation, infinite-scroll pagination, and date-based filtering for logistics delivery management screens.",
        ],
        "projects": [
            (
                "Connect App",
                "Business Networking Application (React Native)",
                "Built a business networking app enabling companies to create and share professional "
                "profiles, manage connection requests, and handle subscriptions. Implemented deep "
                "linking, profile sharing via shortened URLs, group navigation, and username auto-sync.",
            ),
            (
                "FUNZCART / CloudPOS Web",
                "Multi-Tenant E-Commerce Platform (Next.js)",
                "Built and maintained a multi-tenant Next.js platform with 8+ production theme packs "
                "and tenant storefronts. Designed dynamic property filter architecture, live search "
                "autocomplete, and a review system with multi-image uploads. Integrated Razorpay "
                "payments and WhatsApp order notifications.",
            ),
            (
                "Juice World",
                "Production Storefront (Next.js · i18n · ISR)",
                "Developed a production storefront with API-driven content. Implemented English/Arabic "
                "i18n with full RTL support, ISR caching, dynamic SEO metadata, and sitemap generation.",
            ),
            (
                "Ganvin",
                "Logistics & Delivery Ecosystem (Flutter)",
                "Built Executive and Customer apps for the Ganvin logistics ecosystem. Implemented "
                "Google Maps navigation, infinite-scroll pagination, Razorpay payments, and resolved "
                "timezone/API issues with the backend team.",
            ),
            (
                "EPOSMOB",
                "Mobile POS & Inventory Management (Flutter)",
                "Developed a mobile POS application covering billing, Day Close workflows, bilingual "
                "receipt printing, stock reports, purchase orders, and supplier/customer management.",
            ),
        ],
        "skills": [
            ("Mobile Development", "React Native, Deep Linking, Flutter, Dart"),
            ("React Ecosystem", "React.js, Next.js, TypeScript, JavaScript"),
            ("Web Architecture", "SSR, ISR, Dynamic Routing, SEO, i18n, RTL Support, Responsive Design"),
            ("Frontend", "HTML5, CSS3, Tailwind CSS, State Management"),
            ("Integrations", "Razorpay, Google Maps, Algolia Search, WhatsApp Notifications, Barcode Generation"),
            ("Commerce & Multi-Tenant", "Multi-Tenant Architecture (8+ themes), Tenant Onboarding, Product Search, Checkout Flows"),
            ("Also experienced in", "Python, Django, MySQL, SQLite"),
            ("Tools", "Git, GitHub, VS Code, Postman, Figma"),
        ],
    },
    "frontend": {
        "file": "muhammed-mubashir-k-resume-frontend.docx",
        "subtitle": "Frontend Developer  |  React.js · Next.js · TypeScript · Tailwind CSS",
        "summary": (
            "Frontend Developer with production experience building responsive, performant web "
            "applications using React.js, Next.js, and TypeScript. Delivered 8+ theme packs for a "
            "multi-tenant e-commerce platform, bilingual (English/Arabic) storefronts with RTL support, "
            "and complex search, filtering, and checkout UIs. Skilled in SSR/ISR optimization, "
            "component architecture, and responsive design across devices. Also experienced in "
            "React Native and Flutter for mobile development."
        ),
        "current_bullets": [
            "Engineered a multi-tenant Next.js e-commerce platform with 8+ production theme packs, reusable component architecture, and tenant-specific UI experiences across live client storefronts.",
            "Built live search autocomplete with debounced requests and request cancellation, integrated across multiple theme navigation bars with Algolia-powered product discovery.",
            "Delivered bilingual (English/Arabic) storefronts with full RTL layout support, i18n translation systems, ISR caching, and dynamic SEO metadata with sitemap generation.",
            "Designed category-driven dynamic property filter architecture with multi-select filtering, responsive UI, and a review system supporting multi-image uploads.",
            "Integrated Razorpay payment gateway across tenant checkouts; built OTP authentication, guest checkout, and COD flows with WhatsApp order notifications.",
            "Also developed mobile applications using React Native and Flutter for business networking, POS, and logistics domains.",
        ],
        "trainee_bullets": [
            "Integrated REST APIs for authentication, product catalogs, orders, and checkout across Next.js storefronts with responsive, theme-consistent UI.",
            "Built reusable UI components shared across 8+ theme packs with consistent styling, spacing, and responsive behavior.",
            "Implemented Google Maps integration, infinite-scroll pagination, and date-based filtering with responsive layouts for mobile screens.",
        ],
        "projects": [
            (
                "FUNZCART / CloudPOS Web",
                "Multi-Tenant E-Commerce Platform (Next.js · Tailwind)",
                "Built and maintained a multi-tenant Next.js platform with 8+ production theme packs "
                "(Aurora, Nova, Toy, Silk, Ornament, Harvest) and tenant storefronts including Luzine "
                "Bakes. Designed dynamic property filter architecture, live search autocomplete, and a "
                "review system with multi-image uploads. Built responsive layouts across all themes.",
            ),
            (
                "Juice World",
                "Production Storefront (Next.js · i18n · ISR)",
                "Developed a production storefront with API-driven banners, products, services, and "
                "testimonials. Implemented English/Arabic i18n with full RTL support, ISR caching, "
                "dynamic SEO metadata, and responsive design across all devices.",
            ),
            (
                "Connect App",
                "Business Networking Application (React Native)",
                "Built a business networking app enabling companies to create and share professional "
                "profiles, manage connection requests, and handle subscriptions. Implemented deep "
                "linking, profile sharing via shortened URLs, and group navigation.",
            ),
            (
                "Ganvin",
                "Logistics & Delivery Ecosystem (Flutter)",
                "Built Executive and Customer apps with Google Maps navigation, infinite-scroll "
                "pagination, Razorpay payments, and responsive mobile layouts.",
            ),
            (
                "EPOSMOB",
                "Mobile POS & Inventory Management (Flutter)",
                "Developed a mobile POS application covering billing, Day Close workflows, bilingual "
                "receipt printing, stock reports, and supplier/customer management.",
            ),
        ],
        "skills": [
            ("Frontend", "React.js, Next.js, TypeScript, JavaScript, HTML5, CSS3, Tailwind CSS"),
            ("Web Architecture", "SSR, ISR, Dynamic Routing, SEO, Sitemap Generation, i18n, RTL Support, Responsive Design, State Management"),
            ("UI Engineering", "Component Architecture, Multi-Theme Systems, Reusable Components, Responsive Layouts"),
            ("Integrations", "Razorpay, Algolia Search, Google Maps, WhatsApp Notifications, Barcode Generation"),
            ("Commerce & Multi-Tenant", "Multi-Tenant Architecture (8+ themes), Tenant Onboarding, Product Search, Checkout Flows, Order Management"),
            ("Also experienced in", "React Native, Flutter, Dart, Python, Django, MySQL"),
            ("Tools", "Git, GitHub, VS Code, Postman, Figma"),
        ],
    },
}


def months_since(start: date) -> int:
    today = date.today()
    months = (today.year - start.year) * 12 + (today.month - start.month)
    return max(months, 0)


def set_run(run, size=10, bold=False, color=INK):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color


def paragraph_border_bottom(paragraph, color=BORDER, size="6"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_hyperlink(paragraph, text, url, size=9.5, color=LINK):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), str(color))
    r_pr.append(color_el)
    size_el = OxmlElement("w:sz")
    size_el.set(qn("w:val"), str(int(size * 2)))
    r_pr.append(size_el)
    font_el = OxmlElement("w:rFonts")
    font_el.set(qn("w:ascii"), "Calibri")
    font_el.set(qn("w:hAnsi"), "Calibri")
    r_pr.append(font_el)
    new_run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    new_run.append(text_el)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_section_heading(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(5)
    paragraph_border_bottom(p)
    run = p.add_run(title.upper())
    set_run(run, size=11, bold=True, color=HEADING)
    return p


def add_body_paragraph(doc, text, after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run(run, size=10, color=INK)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.first_line_indent = Inches(-0.2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.12
    run = p.add_run(text)
    set_run(run, size=9.5, color=INK)
    return p


def add_role_header(doc, title, meta):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(2)
    left = p.add_run(title)
    set_run(left, size=10.5, bold=True, color=INK)
    sep = p.add_run("  |  ")
    set_run(sep, size=10, color=MUTED)
    right = p.add_run(meta)
    set_run(right, size=9.5, color=MUTED)
    return p


def build(variant_key):
    v = VARIANTS[variant_key]
    out = PUBLIC / v["file"]

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    styles["Normal"]._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    styles["Normal"].font.size = Pt(10)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(1)
    name = title.add_run("MUHAMMED MUBASHIR K")
    set_run(name, size=20, bold=True, color=HEADING)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(4)
    sub = subtitle.add_run(v["subtitle"])
    set_run(sub, size=10, bold=False, color=MUTED)

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_after = Pt(8)
    set_run(contact.add_run("Kerala, India  |  "), size=9.5, color=MUTED)
    add_hyperlink(contact, "muhammedmubashir720@gmail.com", "mailto:muhammedmubashir720@gmail.com")
    set_run(contact.add_run("  |  +91 8089433955  |  "), size=9.5, color=MUTED)
    add_hyperlink(contact, "github.com/MuhammedMubashir-dev", "https://github.com/MuhammedMubashir-dev")
    set_run(contact.add_run("  |  "), size=9.5, color=MUTED)
    add_hyperlink(contact, "linkedin.com/in/muhammed-mubashir-k", "https://www.linkedin.com/in/muhammed-mubashir-k")
    set_run(contact.add_run("  |  "), size=9.5, color=MUTED)
    add_hyperlink(contact, "muhammed-mubashir-portfolio.netlify.app", "https://muhammed-mubashir-portfolio.netlify.app")

    add_section_heading(doc, "Professional Summary")
    add_body_paragraph(doc, v["summary"])

    add_section_heading(doc, "Experience")
    total_months = months_since(ROLE_START)
    total_dur = f"Apr 2026 – Present · {total_months} mo{'s' if total_months != 1 else ''}"

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(1)
    company = p.add_run("ENKE Consulting Services LLP")
    set_run(company, size=11, bold=True, color=INK)
    sep = p.add_run(f"  |  {total_dur}")
    set_run(sep, size=9.5, color=MUTED)

    curr_months = months_since(PROMOTION_DATE)
    add_role_header(doc, "Jr Application Developer", f"Jul 2026 – Present · {curr_months} mo{'s' if curr_months != 1 else ''}")
    for item in v["current_bullets"]:
        add_bullet(doc, item)

    add_role_header(doc, "Full Stack Developer Trainee", "Apr 2026 – Jul 2026 · 3 mos")
    for item in v["trainee_bullets"]:
        add_bullet(doc, item)

    add_section_heading(doc, "Selected Projects")
    for name, kind, detail in v["projects"]:
        add_role_header(doc, name, kind)
        add_body_paragraph(doc, detail, after=3)

    add_section_heading(doc, "Technical Skills")
    for label, value in v["skills"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        label_run = p.add_run(f"{label}: ")
        set_run(label_run, size=9.5, bold=True, color=INK)
        value_run = p.add_run(value)
        set_run(value_run, size=9.5, color=MUTED)

    add_section_heading(doc, "Impact Snapshot")
    add_body_paragraph(
        doc,
        "7+ Production Applications  |  8+ Production Theme Packs  |  "
        "3+ Bilingual (EN/AR) Storefronts with RTL  |  Multi-Tenant Platform Architecture  |  "
        "Promoted in 3 Months",
        after=0,
    )

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    print(out)


if __name__ == "__main__":
    variant = sys.argv[1] if len(sys.argv) > 1 else "react"
    if variant == "all":
        for key in VARIANTS:
            build(key)
    elif variant in VARIANTS:
        build(variant)
    else:
        print(f"Unknown variant: {variant}")
        print(f"Available: {', '.join(VARIANTS.keys())}, all")
        sys.exit(1)
